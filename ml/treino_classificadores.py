from __future__ import annotations

from dataclasses import dataclass
import os
from pathlib import Path
import shutil
import sys
import time
from typing import Dict, List, Optional, Tuple

import joblib
from docx import Document
from dotenv import find_dotenv, load_dotenv
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import MultinomialNB
from sklearn.pipeline import Pipeline
from sklearn.svm import LinearSVC


RANDOM_STATE = 42
load_dotenv(find_dotenv())


@dataclass
class ResultadoModelo:
    nome: str
    acuracia: float
    relatorio: str
    matriz: List[List[int]]


@dataclass
class MetricasApi:
    amostras: int
    tempo_medio_s: Optional[float]
    taxa_erro_percent: Optional[float]
    observacao: str


def gerar_base_sintetica() -> Tuple[List[str], List[str]]:
    generos = [
        "acao",
        "drama",
        "comedia",
        "ficcao cientifica",
        "terror",
        "romance",
        "suspense",
        "fantasia",
    ]

    base_filme = [
        "me recomende filmes de {g}",
        "quero um filme de {g}",
        "indique 5 filmes de {g}",
        "sugira um longa de {g}",
        "me passa filmes parecidos com interestelar de {g}",
    ]
    base_serie = [
        "me recomende series de {g}",
        "quero uma serie de {g}",
        "indique 3 series de {g}",
        "sugira uma serie com temporadas de {g}",
        "me passa series parecidas com dark de {g}",
    ]
    base_livro = [
        "me recomende livros de {g}",
        "quero um livro de {g}",
        "indique 4 livros de {g}",
        "sugira um romance de {g}",
        "me passa livros parecidos com dune de {g}",
    ]
    base_misto = [
        "me recomende filmes e series de {g}",
        "quero livros e filmes de {g}",
        "indique series e livros de {g}",
        "sugira filmes series e livros de {g}",
        "me passe 10 recomendacoes entre filmes e livros de {g}",
    ]

    textos: List[str] = []
    labels: List[str] = []
    blocos = [
        ("filme", base_filme),
        ("serie", base_serie),
        ("livro", base_livro),
        ("misto", base_misto),
    ]

    for label, templates in blocos:
        for genero in generos:
            for template in templates:
                textos.append(template.format(g=genero))
                labels.append(label)

    # Frases extras para variacao e robustez.
    extras = [
        ("filme", "me indica cinema nacional"),
        ("filme", "recomende longas premiados"),
        ("serie", "quero algo episodico para maratonar"),
        ("serie", "indique series curtas"),
        ("livro", "me recomenda leitura para iniciantes"),
        ("livro", "quero romances classicos"),
        ("misto", "mistura filmes series e leitura no mesmo pacote"),
        ("misto", "quero opcoes de filmes e livros"),
    ]
    for label, frase in extras:
        textos.append(frase)
        labels.append(label)

    return textos, labels


def criar_pipelines() -> Dict[str, Pipeline]:
    vetorizador = TfidfVectorizer(ngram_range=(1, 2), lowercase=True, strip_accents="unicode")
    return {
        "logistic_regression": Pipeline(
            [
                ("tfidf", vetorizador),
                ("clf", LogisticRegression(max_iter=2000, random_state=RANDOM_STATE)),
            ]
        ),
        "linear_svc": Pipeline(
            [
                ("tfidf", vetorizador),
                ("clf", LinearSVC(random_state=RANDOM_STATE)),
            ]
        ),
        "multinomial_nb": Pipeline(
            [
                ("tfidf", vetorizador),
                ("clf", MultinomialNB()),
            ]
        ),
    }


def treinar_avaliar() -> Tuple[Dict[str, ResultadoModelo], str, Pipeline, List[str]]:
    textos, labels = gerar_base_sintetica()

    x_train, x_test, y_train, y_test = train_test_split(
        textos,
        labels,
        test_size=0.25,
        random_state=RANDOM_STATE,
        stratify=labels,
    )

    pipelines = criar_pipelines()
    resultados: Dict[str, ResultadoModelo] = {}
    melhor_nome = ""
    melhor_pipeline: Pipeline | None = None
    melhor_acc = -1.0

    for nome, pipeline in pipelines.items():
        pipeline.fit(x_train, y_train)
        pred = pipeline.predict(x_test)
        acc = accuracy_score(y_test, pred)
        relatorio = classification_report(y_test, pred, digits=4)
        matriz = confusion_matrix(y_test, pred, labels=["filme", "serie", "livro", "misto"]).tolist()

        resultados[nome] = ResultadoModelo(nome=nome, acuracia=acc, relatorio=relatorio, matriz=matriz)

        if acc > melhor_acc:
            melhor_acc = acc
            melhor_nome = nome
            melhor_pipeline = pipeline

    if melhor_pipeline is None:
        raise RuntimeError("Nao foi possivel treinar nenhum modelo.")

    return resultados, melhor_nome, melhor_pipeline, y_test


def avaliar_api_real(amostras: int = 6) -> MetricasApi:
    """
    Mede tempo medio e taxa de erro usando chamadas reais ao /api/chat.
    Nao usa mock.
    """
    chave = os.getenv("GEMINI_API_KEY", "").strip()
    if not chave or chave == "COLE_SUA_CHAVE_AQUI":
        return MetricasApi(
            amostras=0,
            tempo_medio_s=None,
            taxa_erro_percent=None,
            observacao="Nao medido: GEMINI_API_KEY ausente ou placeholder no .env.",
        )

    raiz_projeto = Path(__file__).resolve().parents[1]
    if str(raiz_projeto) not in sys.path:
        sys.path.insert(0, str(raiz_projeto))

    from backend import app as backend_app_module

    if backend_app_module.model is None:
        return MetricasApi(
            amostras=0,
            tempo_medio_s=None,
            taxa_erro_percent=None,
            observacao="Nao medido: modelo Gemini nao foi inicializado no backend.",
        )

    mensagens = [
        "Me recomende 5 filmes de suspense",
        "Quero 4 series parecidas com Dark",
        "Me indique 3 livros de fantasia",
        "Quero filmes e livros de ficcao cientifica",
        "Me recomenda 2 series de comedia",
        "Quero 6 filmes de acao",
    ][: max(1, amostras)]

    tempos: List[float] = []
    erros = 0
    backend_app_module._ip_hits.clear()

    original_testing = backend_app_module.app.config.get("TESTING", False)
    backend_app_module.app.config["TESTING"] = False
    try:
        with backend_app_module.app.test_client() as client:
            for msg in mensagens:
                inicio = time.perf_counter()
                resposta = client.post("/api/chat", json={"message": msg, "stream": False})
                fim = time.perf_counter()
                tempos.append(fim - inicio)
                if resposta.status_code >= 400:
                    erros += 1
    finally:
        backend_app_module.app.config["TESTING"] = original_testing

    total = len(mensagens)
    if total == 0:
        return MetricasApi(
            amostras=0,
            tempo_medio_s=None,
            taxa_erro_percent=None,
            observacao="Nao medido: sem amostras disponiveis.",
        )

    tempo_medio = sum(tempos) / total
    taxa_erro = (erros / total) * 100.0
    return MetricasApi(
        amostras=total,
        tempo_medio_s=tempo_medio,
        taxa_erro_percent=taxa_erro,
        observacao="Metrica coletada em execucao real do backend local (/api/chat), sem mock.",
    )


def copiar_relatorio_para_downloads(relatorio_path: Path) -> Tuple[Optional[Path], str]:
    downloads = Path.home() / "Downloads"
    if not downloads.exists():
        return None, "Downloads nao encontrado; copia automatica nao realizada."
    destino = downloads / "RELATORIO_N2.docx"
    try:
        shutil.copy2(relatorio_path, destino)
        return destino, "ok"
    except PermissionError:
        return (
            None,
            "Nao foi possivel atualizar Downloads/RELATORIO_N2.docx porque o arquivo esta aberto. "
            "Feche o arquivo e rode novamente.",
        )


def gerar_relatorio_word(
    resultados: Dict[str, ResultadoModelo],
    melhor_nome: str,
    melhor_matriz: List[List[int]],
    total_amostras: int,
    metrica_semantica: float,
    metricas_api: MetricasApi,
    output_path: Path,
) -> None:
    doc = Document()
    labels = ["filme", "serie", "livro", "misto"]

    doc.add_heading("Relatorio N2 - Classificacao de pedidos no EratoChat", level=1)

    doc.add_heading("Objetivo", level=2)
    doc.add_paragraph(
        "Dar continuidade ao EratoChat com um modulo de classificacao supervisionada para "
        "identificar o tipo de obra solicitado pelo usuario (filme, serie, livro ou misto)."
    )

    doc.add_heading("Metodologia", level=2)
    doc.add_paragraph("Base sintetica rotulada com variacoes de pedidos.", style="List Bullet")
    doc.add_paragraph("Vetorizacao de texto com TF-IDF (unigramas e bigramas).", style="List Bullet")
    doc.add_paragraph(
        "Separacao treino/teste com train_test_split (75% treino, 25% teste, estratificado).",
        style="List Bullet",
    )
    doc.add_paragraph("Treino e comparacao de 3 classificadores.", style="List Bullet")

    doc.add_heading("Modelos avaliados", level=2)
    doc.add_paragraph("Logistic Regression", style="List Bullet")
    doc.add_paragraph("Linear SVC", style="List Bullet")
    doc.add_paragraph("Multinomial Naive Bayes", style="List Bullet")

    doc.add_heading("Resultado comparativo (acuracia)", level=2)
    tabela_resultados = doc.add_table(rows=1, cols=2)
    tabela_resultados.style = "Table Grid"
    cabecalho = tabela_resultados.rows[0].cells
    cabecalho[0].text = "Modelo"
    cabecalho[1].text = "Acuracia"

    for nome, resultado in resultados.items():
        row = tabela_resultados.add_row().cells
        row[0].text = nome
        row[1].text = f"{resultado.acuracia:.4f}"

    doc.add_paragraph(f"Melhor modelo: {melhor_nome}.")

    doc.add_heading("Matriz de confusao (melhor modelo)", level=2)
    tabela_matriz = doc.add_table(rows=len(labels) + 1, cols=len(labels) + 1)
    tabela_matriz.style = "Table Grid"
    tabela_matriz.cell(0, 0).text = "real/pred"
    for idx, label in enumerate(labels, start=1):
        tabela_matriz.cell(0, idx).text = label
        tabela_matriz.cell(idx, 0).text = label

    for i, linha in enumerate(melhor_matriz, start=1):
        for j, valor in enumerate(linha, start=1):
            tabela_matriz.cell(i, j).text = str(valor)

    doc.add_heading("Analise da matriz e aderencia ao objetivo", level=2)
    doc.add_paragraph(
        "A diagonal principal representa os acertos por classe e deve concentrar a maior parte dos valores.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Erros entre classes semelhantes (ex.: filme vs misto) sao esperados em pedidos curtos ou ambiguos.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "O uso da classe misto ajuda a preservar a logica atual do projeto quando o usuario pede mais de um tipo.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Este classificador pode ser acoplado ao backend para reforcar validacoes de tipo antes do prompt final.",
        style="List Bullet",
    )

    doc.add_paragraph(f"Total de amostras usadas: {total_amostras}.", style="List Bullet")

    doc.add_heading("Metricas de sucesso (medicao real)", level=2)
    doc.add_paragraph(
        f"Precisao semantica (teste supervisionado): {metrica_semantica:.2f}% "
        f"{'(atingiu >= 85%)' if metrica_semantica >= 85 else '(abaixo de 85%)'}",
        style="List Bullet",
    )
    if metricas_api.tempo_medio_s is None or metricas_api.taxa_erro_percent is None:
        doc.add_paragraph("Tempo de resposta: nao medido.", style="List Bullet")
        doc.add_paragraph("Taxa de erro da API: nao medida.", style="List Bullet")
    else:
        doc.add_paragraph(
            f"Tempo medio de resposta API: {metricas_api.tempo_medio_s:.2f}s "
            f"{'(atingiu < 5s)' if metricas_api.tempo_medio_s < 5 else '(acima de 5s)'}",
            style="List Bullet",
        )
        doc.add_paragraph(
            f"Taxa de erro da API: {metricas_api.taxa_erro_percent:.2f}% "
            f"{'(atingiu < 3%)' if metricas_api.taxa_erro_percent < 3 else '(acima de 3%)'}",
            style="List Bullet",
        )
    doc.add_paragraph(metricas_api.observacao)

    doc.add_heading("Escalabilidade (estado atual)", level=2)
    doc.add_paragraph(
        "Backend stateless e compativel com escalabilidade horizontal no deploy atual.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Itens de evolucao ainda nao implementados: cache de respostas, banco de dados dedicado e load balancer.",
        style="List Bullet",
    )

    doc.save(output_path)


def main() -> None:
    raiz = Path(__file__).resolve().parent
    artefatos = raiz / "artefatos"
    artefatos.mkdir(exist_ok=True)

    resultados, melhor_nome, melhor_pipeline, _ = treinar_avaliar()
    total_amostras = len(gerar_base_sintetica()[0])
    melhor_matriz = resultados[melhor_nome].matriz
    metrica_semantica = resultados[melhor_nome].acuracia * 100.0
    metricas_api = avaliar_api_real(amostras=6)

    joblib.dump(melhor_pipeline, artefatos / "modelo_tipo_obra.joblib")

    metricas_path = artefatos / "metricas_modelos.txt"
    with metricas_path.open("w", encoding="utf-8") as f:
        for nome, resultado in resultados.items():
            f.write(f"=== {nome} ===\n")
            f.write(f"acuracia: {resultado.acuracia:.4f}\n")
            f.write("classification_report:\n")
            f.write(resultado.relatorio)
            f.write("\n\n")
        f.write("=== metricas_sucesso ===\n")
        f.write(f"precisao_semantica_percent: {metrica_semantica:.2f}\n")
        if metricas_api.tempo_medio_s is not None:
            f.write(f"tempo_medio_api_segundos: {metricas_api.tempo_medio_s:.4f}\n")
            f.write(f"taxa_erro_api_percent: {metricas_api.taxa_erro_percent:.2f}\n")
        else:
            f.write("tempo_medio_api_segundos: NA\n")
            f.write("taxa_erro_api_percent: NA\n")
        f.write(f"observacao_api: {metricas_api.observacao}\n")

    relatorio_path = raiz / "RELATORIO_N2.docx"
    gerar_relatorio_word(
        resultados,
        melhor_nome,
        melhor_matriz,
        total_amostras,
        metrica_semantica,
        metricas_api,
        relatorio_path,
    )
    destino_downloads, status_downloads = copiar_relatorio_para_downloads(relatorio_path)

    print("Treino concluido com sucesso.")
    print(f"Melhor modelo: {melhor_nome} ({resultados[melhor_nome].acuracia:.4f})")
    print(f"Artefatos em: {artefatos}")
    print(f"Relatorio: {relatorio_path}")
    if destino_downloads is not None:
        print(f"Relatorio atualizado em Downloads: {destino_downloads}")
    else:
        print(status_downloads)


if __name__ == "__main__":
    main()
